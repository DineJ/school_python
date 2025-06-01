import tkinter as tk
from tkinter import messagebox, ttk, font, colorchooser, simpledialog
import requests
import sqlite3
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import unittest
import os
import sys
import platform
import subprocess

canvas = None
# Connexion à la base
try:
    conn = sqlite3.connect('data.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS data (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT,
                    state TEXT,
                    size INTEGER,
                    length REAL
                )''')
    conn.commit()
except sqlite3.Error as e:
    messagebox.showerror("Erreur base de données", f"Erreur lors de la connexion : {e}")

# Fonctions principales
def fetch_data():
    try:
        c.execute("SELECT COUNT(*) FROM data")
        if c.fetchone()[0] > 0:
            if not messagebox.askyesno("Confirmation", "La base de données n'est pas vide. Effacer et continuer ?"):
                return
            clear_db()

        url = "https://jsonplaceholder.typicode.com/users"
        response = requests.get(url)
        response.raise_for_status()
        data = response.json()

        for item in data:
            name = item.get('name')
            address = item.get('address', {})
            state = address.get('city')
            email = item.get('email')

            if not (name and state and email):
                # On ignore l'entrée si une donnée est manquante
                continue

            size = len(name)
            length = len(email)
            c.execute("INSERT INTO data (name, state, size, length) VALUES (?, ?, ?, ?)", (name, state, size, length))
        conn.commit()
        status_var.set("Données téléchargées et stockées avec succès")
    except requests.RequestException as e:
        messagebox.showerror("Erreur", f"Erreur de téléchargement : {e}")
    except sqlite3.Error as e:
        messagebox.showerror("Erreur base de données", str(e))


def clear_db():
    global canvas
    try:
        # Détruire le canvas s'il existe déjà
        if canvas is not None:
            canvas.get_tk_widget().destroy()
        result_label.config(text=f"")
        c.execute("DELETE FROM data")
        conn.commit()
        status_var.set("Base de données effacée")
    except sqlite3.Error as e:
        messagebox.showerror("Erreur", str(e))

def aggregate_data():
    try:
        c.execute("SELECT AVG(length) FROM data")
        avg = c.fetchone()[0]
        result_label.config(text=f"Longueur moyenne des adresses mails : {avg:.2f}" if avg else "Aucune donnée")
        if not avg:
            status_var.set("Aucune donnée")
        else:
            status_var.set("Agrégat calculé")
    except sqlite3.Error as e:
        messagebox.showerror("Erreur", str(e))

def show_graph():
    global canvas
    try:
        # Détruire le canvas s'il existe déjà
        if canvas is not None:
            canvas.get_tk_widget().destroy()

        c.execute("SELECT AVG(length) FROM data")
        avg = c.fetchone()[0]
        if not avg:
            result_label.config(text=f"Aucune donnée")
            return

        c.execute("SELECT size FROM data")
        sizes = [row[0] for row in c.fetchall()]

        fig, ax = plt.subplots()
        ax.hist(sizes, bins=5)
        ax.set_title("Distribution des tailles de noms")
        ax.set_xlabel("Taille")
        ax.set_ylabel("Fréquence")

        canvas = FigureCanvasTkAgg(fig, master=main_frame)
        canvas.draw()
        canvas.get_tk_widget().pack()
        status_var.set("Graphique affiché")
    except sqlite3.Error as e:
        messagebox.showerror("Erreur", str(e))


# Fonctions options
def change_bg_color():
    try:
        color = colorchooser.askcolor()[1]
        if color:
            root.configure(bg=color)
            main_frame.configure(style="Custom.TFrame")
            style.configure("Custom.TFrame", background=color)
            status_var.set(f"Couleur de fond changée en {color}")
    except Exception as e:
        messagebox.showerror("Erreur", f"Impossible de changer la couleur : {e}")


def change_font():
    try:
        size = simpledialog.askinteger("Taille police", "Entrez la taille de la police (ex : 10-30):", minvalue=8, maxvalue=72)
        if size:
            new_font = font.Font(family="Helvetica", size=size, weight="bold")
            result_label.config(font=new_font)
            status_var.set(f"Taille de police changée en {size}")
    except Exception as e:
        messagebox.showerror("Erreur", f"Impossible de changer la police : {e}")


# === Partie 2 intégrée ===

def lancer_analyse_livre():
    try:
        import requests
        from PIL import Image, ImageOps
        import io
        from docx import Document
        from docx.shared import Inches
        import matplotlib.pyplot as plt
        from collections import Counter
        import re
        import os

        url = "https://www.gutenberg.org/files/11/11-0.txt"
        text = requests.get(url).text

        lines = text.splitlines()
        title = next((line for line in lines if line.strip().lower().startswith("alice")), "Titre inconnu")
        author = next((line for line in lines if "lewis carroll" in line.lower()), "Auteur inconnu")

        matches = list(re.finditer(r"CHAPTER I\.", text))
        if len(matches) >= 2:
            start_index = matches[1].end()
            text = text[start_index:]
            end_chapter = re.search(r"CHAPTER II\.", text)
            first_chapter_text = text[:end_chapter.start()] if end_chapter else ""
        else:
            first_chapter_text = ""

        first_chapter_text = first_chapter_text.replace("Down the Rabbit-Hole", "")
        paragraphs = [p.strip() for p in first_chapter_text.split("\n\n") if p.strip()]
        word_counts = [len(p.split()) for p in paragraphs]

        def round_to_tens(wc):
            return wc // 10 * 10 if wc >= 10 else None

        rounded_counts = [round_to_tens(wc) for wc in word_counts]
        valid_paragraphs = [p for p, wc in zip(paragraphs, rounded_counts) if wc is not None]
        valid_word_counts = [wc for wc in rounded_counts if wc is not None]
        count_freq = Counter(valid_word_counts)

        plt.figure(figsize=(8, 6))
        plt.bar(count_freq.keys(), count_freq.values(), width=8, edgecolor='black')
        plt.xlabel("Nombre de mots (arrondi à la dizaine)")
        plt.ylabel("Nombre de paragraphes")
        plt.title("Distribution des longueurs des paragraphes du premier chapitre")
        plt.tight_layout()
        plt.savefig("distribution_paragraphes.png")
        plt.close()

        img1 = None
        try:
            img_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/6/6b/Front_cover_Alice_in_Wonderland_%281917%29_01.jpg/640px-Front_cover_Alice_in_Wonderland_%281917%29_01.jpg"
            headers = { "User-Agent": "Mozilla/5.0" }
            response = requests.get(img_url, headers=headers)
            if response.status_code == 200:
                img1 = Image.open(io.BytesIO(response.content)).convert("RGBA")
                img1 = img1.resize((300, 300))
                img1 = ImageOps.fit(img1, (300, 300))
        except Exception as e:
            print(f"Erreur téléchargement image 1 : {str(e)}")

        try:
            logo = Image.open("ynov_logo.png").convert("L")
            threshold = 200
            logo_rgba = Image.new("RGBA", logo.size)
            datas = [(0, 0, 0, 255 if p < threshold else 0) for p in logo.getdata()]
            logo_rgba.putdata(datas)
            rotated_logo = logo_rgba.rotate(30, expand=True)
            img1.paste(rotated_logo, (0, 0), rotated_logo if rotated_logo.mode == 'RGBA' else None)
        except Exception as e:
            print(f"Erreur traitement logo : {str(e)}")

        doc = Document()
        doc.add_heading("Rapport de lecture", 0)
        doc.add_heading(title, level=1)
        if img1:
            temp_file = "photo1.png"
            img1.save(temp_file)
            doc.add_picture(temp_file, width=Inches(4))

        doc.add_paragraph(f"Auteur du livre : {author}")
        doc.add_paragraph("Auteur du rapport : Jridi Dine")
        doc.add_page_break()
        doc.add_heading("Graphique de distribution", level=1)
        doc.add_picture("distribution_paragraphes.png", width=Inches(5))
        doc.add_paragraph(
            f"Nombre de paragraphes : {len(valid_paragraphs)}\n"
            f"Total mots : {sum(valid_word_counts)}\n"
            f"Min : {min(valid_word_counts)}, Max : {max(valid_word_counts)}\n"
            f"Moyenne : {sum(valid_word_counts)//len(valid_word_counts)}\n"
            f"Source : Project Gutenberg"
        )

        doc.save("rapport_final.docx")
        try:
            # Ouvrir automatiquement le fichier
            system = platform.system()
            if system == "Windows":
                os.startfile("rapport_final.docx")
            elif system == "Darwin":  # macOS
                subprocess.run(["open", "rapport_final.docx"])
            else:  # Linux
                subprocess.run(["xdg-open", "rapport_final.docx"])
            messagebox.showinfo("Rapport", "Rapport généré avec succès.")
        except Exception as e:
            messagebox.showwarning("Ouverture du fichier", f"Le fichier a été généré, mais n'a pas pu être ouvert automatiquement : {e}")
    except Exception as e:
        messagebox.showerror("Erreur", str(e))

# Interface
root = tk.Tk()
root.title("Application Desktop - Partie 1 & 2")
root.geometry("800x780")

style = ttk.Style()
style.configure("TFrame", padding=10)

menu_bar = tk.Menu(root)

# Menu Fichier
file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="Quitter", command=root.quit)
menu_bar.add_cascade(label="Fichier", menu=file_menu)

analyse_menu = tk.Menu(menu_bar, tearoff=0)
analyse_menu.add_command(label="Analyse", command=lancer_analyse_livre)
menu_bar.add_cascade(label="Livre", menu=analyse_menu)

# Menu Base de données
db_menu = tk.Menu(menu_bar, tearoff=0)
db_menu.add_command(label="Effacer la base", command=clear_db)
db_menu.add_command(label="Télécharger les données", command=fetch_data)
menu_bar.add_cascade(label="Base de données", menu=db_menu)

# Menu Options
options_menu = tk.Menu(menu_bar, tearoff=0)
options_menu.add_command(label="Changer couleur de fond", command=change_bg_color)
options_menu.add_command(label="Changer la police", command=change_font)
menu_bar.add_cascade(label="Options", menu=options_menu)

root.config(menu=menu_bar)

main_frame = ttk.Frame(root)
main_frame.pack(pady=10)

ttk.Button(main_frame, text="Afficher moyenne", command=aggregate_data).pack(pady=5)
ttk.Button(main_frame, text="Afficher graphique", command=show_graph).pack(pady=5)
result_label = ttk.Label(main_frame, text="")
result_label.pack(pady=5)

status_var = tk.StringVar()
status_bar = ttk.Label(root, textvariable=status_var, relief=tk.SUNKEN, anchor="w")
status_bar.pack(side=tk.BOTTOM, fill=tk.X)

def on_closing():
    try:
        clear_db()
    except Exception as e:
        messagebox.showerror("Erreur", f"Erreur lors de la suppression des données : {e}")
    try:
        conn.close()
    except Exception:
        pass
    root.destroy()
    sys.exit(0)

root.protocol("WM_DELETE_WINDOW", on_closing)

root.mainloop()


